"""
带批注的 .docx 导出器 —— 将原始文档 + 5 Agent 评审结果导出为 Word 批注文档。

使用 python-docx + lxml 生成真正的 Word 批注（comment），
包含：A4 语法修正（删除线+加粗）、段落批注、总结表格。

用法:
    from app.utils.docx_exporter import export_annotated_docx
    output_path = export_annotated_docx(review_result, output_dir="/tmp")
"""

from __future__ import annotations

import datetime
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

from app.utils.text_utils import split_paragraphs, parse_location


# ── Word Namespaces ──
WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"


def export_annotated_docx(
    result: dict,
    output_dir: str | None = None,
) -> str:
    """将评审结果导出为 .docx 批注文档。

    Args:
        result: 来自 /api/v1/review 的完整评审 JSON（包含 rubric/structure/argument/language/integrity/meta）
        output_dir: 输出目录（默认系统临时目录）

    Returns:
        生成的文件路径
    """
    if output_dir is None:
        output_dir = tempfile.gettempdir()
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # 提取原始文本
    annotated_md = result.get("meta", {}).get("annotated_md", "")
    original_text = _extract_original_text(annotated_md, result)

    agent_data = {
        "structure": result.get("structure"),
        "argument": result.get("argument"),
        "language": result.get("language"),
        "integrity": result.get("integrity"),
        "rubric": result.get("rubric"),
    }

    # 创建文档
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_heading("审稿标注报告", level=0)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"AcademicReviewer v0.6.0  |  "
        f"竞赛: {result.get('competition', '?')}  |  "
        f"总评分: {result.get('total_score', '?')}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph("")  # spacer

    # 图例
    legend = doc.add_paragraph()
    legend.style = doc.styles["Normal"]
    _add_legend_run(legend, "A2 = 结构逻辑", RGBColor(0x25, 0x63, 0xEB))
    _add_legend_run(legend, "  |  A3 = 论点证据", RGBColor(0x16, 0xA3, 0x4A))
    _add_legend_run(legend, "  |  A4 = 语言风格", RGBColor(0xEA, 0x58, 0x0C))
    _add_legend_run(legend, "  |  A5 = 学术诚信", RGBColor(0xDC, 0x26, 0x26))
    legend.paragraph_format.space_after = Pt(6)

    conf_legend = doc.add_paragraph()
    _add_legend_run(conf_legend, "⚡ = 需老师确认(REVIEW)", RGBColor(0x92, 0x4E, 0x0E))
    _add_legend_run(conf_legend, "  |  🔴 = 必须人工审核(ESCALATE)", RGBColor(0xDC, 0x26, 0x26))
    _add_legend_run(conf_legend, "  |  无标记 = AI可信(FULL)", RGBColor(0x16, 0xA3, 0x4A))
    conf_legend.paragraph_format.space_after = Pt(18)

    # ── 逐段渲染 ──
    paragraphs = split_paragraphs(original_text)
    if not paragraphs:
        p = doc.add_paragraph(original_text or "(空白文档)")
        p.style = doc.styles["Normal"]

    # 建立注解索引：段落索引 → [注解列表]
    annotations_by_para: dict[int, list[dict]] = {i: [] for i in range(len(paragraphs))}

    # A4 rewrites
    lang = agent_data.get("language") or {}
    for rw in lang.get("rewrites", []):
        idx = parse_location(rw.get("location", ""))
        if 0 <= idx < len(paragraphs):
            annotations_by_para[idx].append({
                "type": "rewrite",
                "agent": "A4",
                "issue": rw.get("issue", "?"),
                "original": rw.get("original", ""),
                "corrected": rw.get("corrected", ""),
                "substitutability": rw.get("substitutability", ""),
            })

    # A4 suggestions
    for sg in lang.get("suggestions", []):
        idx = parse_location(sg.get("location", ""))
        if 0 <= idx < len(paragraphs):
            annotations_by_para[idx].append({
                "type": "suggestion",
                "agent": "A4",
                "issue": sg.get("type", "?"),
                "description": sg.get("description", ""),
                "hint": sg.get("hint", ""),
                "severity": sg.get("severity", ""),
                "substitutability": sg.get("substitutability", ""),
            })

    # A2 structure issues
    struct = agent_data.get("structure") or {}
    for si in struct.get("section_issues", []):
        section_name = si.get("section", "")
        for i, para in enumerate(paragraphs):
            if section_name.lower() in para.lower()[:120]:
                annotations_by_para[i].append({
                    "type": "section_issue",
                    "agent": "A2",
                    "section": section_name,
                    "issue": si.get("current_problem", si.get("issue", "")),
                    "hint": si.get("hint", ""),
                    "severity": si.get("severity", ""),
                    "substitutability": si.get("substitutability", ""),
                })
                break

    # A3 logical fallacies
    arg = agent_data.get("argument") or {}
    for fl in arg.get("logical_fallacies", []):
        idx = parse_location(fl.get("location", ""))
        if 0 <= idx < len(paragraphs):
            annotations_by_para[idx].append({
                "type": "fallacy",
                "agent": "A3",
                "fallacy_type": fl.get("fallacy_type", "?"),
                "description": fl.get("description", fl.get("why_it_fails", "")),
                "correct_form": fl.get("correct_form", ""),
                "severity": fl.get("severity", ""),
                "substitutability": fl.get("substitutability", ""),
            })

    # ── 渲染每一段 ──
    comment_id = 0  # 递增的批注 ID
    comment_parts: list[etree._Element] = []  # 收集所有批注 XML 元素

    for i, para_text in enumerate(paragraphs):
        # 段落标题
        p_label = doc.add_paragraph()
        p_label.style = doc.styles["Normal"]
        run_label = p_label.add_run(f"【第 {i+1} 段】")
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        # 段落正文
        p_body = doc.add_paragraph()
        p_body.style = doc.styles["Normal"]
        p_body.paragraph_format.line_spacing = 1.5
        run_body = p_body.add_run(para_text)
        run_body.font.size = Pt(11)

        # 该段的注解
        anns = annotations_by_para.get(i, [])
        if not anns:
            doc.add_paragraph("")  # spacer
            continue

        # 在段落后添加批注标记区域
        for ann in anns:
            emoji = _get_emoji(ann.get("substitutability", ""))
            agent = ann["agent"]
            color = _get_agent_color(agent)
            ann_type = ann["type"]

            p_ann = doc.add_paragraph()
            p_ann.paragraph_format.left_indent = Cm(1)
            p_ann.paragraph_format.space_before = Pt(4)
            p_ann.paragraph_format.space_after = Pt(4)

            if ann_type == "rewrite":
                # 修正条目：删除线原文 → 加粗修正
                run_tag = p_ann.add_run(f"{emoji}[{agent}—{ann.get('issue', '?')}] ")
                run_tag.bold = True
                run_tag.font.size = Pt(9)
                run_tag.font.color.rgb = color

                # 原文（删除线）
                run_orig = p_ann.add_run(ann.get("original", ""))
                run_orig.font.strike = True
                run_orig.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                run_orig.font.size = Pt(10)

                # 箭头
                run_arrow = p_ann.add_run("  →  ")
                run_arrow.font.size = Pt(9)
                run_arrow.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

                # 修正（加粗绿色）
                run_corr = p_ann.add_run(ann.get("corrected", ""))
                run_corr.bold = True
                run_corr.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
                run_corr.font.size = Pt(10)

                # 添加 Word 批注
                comment_id += 1
                comment_text = (
                    f"[{agent}] {ann.get('issue', '?')}\n"
                    f"原文: {ann.get('original', '')}\n"
                    f"修正: {ann.get('corrected', '')}"
                )
                _inject_comment_marker(p_body, comment_id)
                comment_parts.append(
                    _build_comment_xml(comment_id, comment_text, agent)
                )

            elif ann_type == "suggestion":
                sev = ann.get("severity", "")
                sev_label = {"high": "‼️", "medium": "⚠️", "low": "💡"}.get(sev, "")
                run_tag = p_ann.add_run(
                    f"{emoji}[{agent}—{ann.get('issue', '?')}] {sev_label} "
                )
                run_tag.bold = True
                run_tag.font.size = Pt(9)
                run_tag.font.color.rgb = color

                run_desc = p_ann.add_run(ann.get("description", ""))
                run_desc.font.size = Pt(10)
                run_desc.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

                if ann.get("hint"):
                    run_hint = p_ann.add_run(f"\n          💡 {ann['hint']}")
                    run_hint.font.size = Pt(9)
                    run_hint.font.italic = True
                    run_hint.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

                comment_id += 1
                comment_text = (
                    f"[{agent}] 风格建议 — {ann.get('issue', '?')} ({sev})\n"
                    f"{ann.get('description', '')}\n"
                )
                if ann.get("hint"):
                    comment_text += f"建议: {ann['hint']}"
                _inject_comment_marker(p_body, comment_id)
                comment_parts.append(
                    _build_comment_xml(comment_id, comment_text, agent)
                )

            elif ann_type in ("section_issue", "fallacy"):
                sev = ann.get("severity", "")
                sev_label = {"high": "‼️", "medium": "⚠️", "low": "💡"}.get(sev, "")
                if ann_type == "section_issue":
                    label = f"{ann.get('section', '?')} 章节问题"
                else:
                    label = f"逻辑谬误: {ann.get('fallacy_type', '?')}"

                run_tag = p_ann.add_run(f"{emoji}[{agent}] {label} {sev_label}")
                run_tag.bold = True
                run_tag.font.size = Pt(9)
                run_tag.font.color.rgb = color

                run_desc = p_ann.add_run(f"\n{ann.get('issue', ann.get('description', ''))}")
                run_desc.font.size = Pt(10)
                run_desc.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

                if ann.get("hint"):
                    run_hint = p_ann.add_run(f"\n          💡 {ann['hint']}")
                    run_hint.font.size = Pt(9)
                    run_hint.font.italic = True
                    run_hint.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

                if ann.get("correct_form"):
                    run_cf = p_ann.add_run(f"\n          正确形式: {ann['correct_form']}")
                    run_cf.font.size = Pt(9)
                    run_cf.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

                comment_id += 1
                comment_text = f"[{agent}] {label} ({sev})\n{ann.get('issue', ann.get('description', ''))}"
                if ann.get("hint"):
                    comment_text += f"\n\n建议: {ann['hint']}"
                _inject_comment_marker(p_body, comment_id)
                comment_parts.append(
                    _build_comment_xml(comment_id, comment_text, agent)
                )

        doc.add_paragraph("")  # paragraph spacer

    # ── 总结章节 ──
    doc.add_page_break()
    doc.add_heading("评审总结", level=1)

    # 分数表格
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "审查员"
    hdr[1].text = "维度"
    hdr[2].text = "评分"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    scores = result.get("scores", {})
    score_rows = [
        ("A1 RubricParser", "评分标准解析", scores.get("RubricParser", "-")),
        ("A2 StructureLogic", "结构与逻辑", scores.get("StructureLogic", "-")),
        ("A3 ArgumentEvidence", "论点与证据", scores.get("ArgumentEvidence", "-")),
        ("A4 LanguageStyle", "语言与风格", scores.get("LanguageStyle", "-")),
        ("A5 AcademicIntegrity", "学术诚信", scores.get("AcademicIntegrity", "-")),
    ]
    for agent, dim, score in score_rows:
        row = table.add_row()
        row.cells[0].text = agent
        row.cells[1].text = dim
        row.cells[2].text = str(score)

    doc.add_paragraph("")

    # 关键发现
    for agent_name, data, section_title in [
        ("结构 (A2)", agent_data.get("structure"), "结构逻辑"),
        ("论证 (A3)", agent_data.get("argument"), "论点证据"),
        ("语言 (A4)", agent_data.get("language"), "语言风格"),
        ("学术诚信 (A5)", agent_data.get("integrity"), "学术诚信"),
    ]:
        if not data or not isinstance(data, dict):
            continue
        score_key = {
            "结构 (A2)": "structure_score",
            "论证 (A3)": "overall_score",
            "语言 (A4)": None,
            "学术诚信 (A5)": "integrity_score",
        }[agent_name]

        doc.add_heading(f"{section_title}", level=2)
        if score_key:
            sc = data.get(score_key, "-")
            doc.add_paragraph(f"评分: {sc}", style="List Bullet")

        for pp in data.get("positive_points", []):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"✅ {pp}")
            run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

        for ki in data.get("key_issues", []):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"⚠️ {ki}")
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    # 校信度统计
    if agent_data.get("language"):
        lang_summary = agent_data["language"].get("summary", {})
        if lang_summary:
            doc.add_heading("语言统计", level=2)
            grammar = lang_summary.get("grammar_score", "-")
            style = lang_summary.get("style_score", "-")
            overall = lang_summary.get("overall_language_score", "-")
            doc.add_paragraph(f"语法: {grammar}  |  风格: {style}  |  综合: {overall}")
            rc = lang_summary.get("rewrite_count", 0)
            sc = lang_summary.get("suggestion_count", 0)
            doc.add_paragraph(f"修正 {rc} 处  |  风格建议 {sc} 条")

    # 元数据
    doc.add_heading("评审元数据", level=2)
    meta = result.get("meta", {})
    meta_items = [
        f"总评分: {result.get('total_score', '?')}",
        f"竞赛: {result.get('competition', '?')}",
        f"竞赛类型: {result.get('competition_type', '?')}",
        f"字数: {meta.get('word_count', '?')}",
        f"模型: {meta.get('model', '?')}",
        f"耗时: {meta.get('duration_seconds', '?')}s",
    ]
    student_level = meta.get("student_level", "")
    if student_level:
        meta_items.append(f"学生水平: {student_level}")
    for item in meta_items:
        doc.add_paragraph(item, style="List Bullet")

    # ── 构建批注 XML ──
    comments_xml = None
    if comment_parts:
        comments_root = etree.Element(f"{{{WML_NS}}}comments")
        for elem in comment_parts:
            comments_root.append(elem)
        comments_xml = etree.tostring(
            comments_root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    # ── 保存 ──
    competition = result.get("competition", "review")
    safe_name = re.sub(r"[^a-zA-Z0-9_一-鿿\-]", "_", str(competition))
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"AcademicReviewer_{safe_name}_{timestamp}.docx"
    output_path = out_dir / filename
    doc.save(str(output_path))

    # ── 保存后将批注注入到 docx zip ──
    if comments_xml:
        _finalize_comments(str(output_path), comments_xml)

    return str(output_path)


# ── Word 批注 XML 底层操作 ──

def _inject_comment_marker(paragraph, comment_id: int) -> None:
    """在段落的 XML 中插入批注范围标记。

    在每个段落的末尾插入 <w:commentRangeStart> 和 <w:commentRangeEnd>
    以及 <w:commentReference>，使 Word 能在该段落显示批注气泡。
    """
    p_element = paragraph._element

    # 批注范围开始
    start = parse_xml(
        f'<w:commentRangeStart {nsdecls("w")} w:id="{comment_id}"/>'
    )
    # 批注范围结束
    end = parse_xml(
        f'<w:commentRangeEnd {nsdecls("w")} w:id="{comment_id}"/>'
    )
    # 批注引用标记
    ref = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
        f'<w:commentReference w:id="{comment_id}"/>'
        f'</w:r>'
    )

    # 插入到段落开头（批注范围覆盖整段）
    first_run = p_element.find(qn("w:r"))
    if first_run is not None:
        p_element.insert(list(p_element).index(first_run), start)
        p_element.insert(list(p_element).index(first_run) + 1, end)
    else:
        p_element.append(start)
        p_element.append(end)

    # 引用标记放在段落末尾
    p_element.append(ref)


def _build_comment_xml(comment_id: int, text: str, author: str) -> etree._Element:
    """构建单个批注的 XML 元素（用于 comments.xml）。"""
    now = datetime.datetime.now().isoformat()

    comment = etree.SubElement(
        etree.Element(f"{{{WML_NS}}}comments"),
        f"{{{WML_NS}}}comment",
    )
    comment.set(f"{{{WML_NS}}}id", str(comment_id))
    comment.set(f"{{{WML_NS}}}author", f"AcademicReviewer-{author}")
    comment.set(f"{{{WML_NS}}}date", now)

    # 将文本按行分割为多个段落
    for line in text.split("\n"):
        p_elem = etree.SubElement(comment, f"{{{WML_NS}}}p")
        p_pr = etree.SubElement(p_elem, f"{{{WML_NS}}}pPr")
        p_style = etree.SubElement(p_pr, f"{{{WML_NS}}}pStyle")
        p_style.set(f"{{{WML_NS}}}val", "CommentText")
        r_elem = etree.SubElement(p_elem, f"{{{WML_NS}}}r")
        r_pr = etree.SubElement(r_elem, f"{{{WML_NS}}}rPr")
        r_fonts = etree.SubElement(r_pr, f"{{{WML_NS}}}rFonts")
        r_fonts.set(f"{{{WML_NS}}}ascii", "Microsoft YaHei")
        r_fonts.set(f"{{{WML_NS}}}hAnsi", "Microsoft YaHei")
        t_elem = etree.SubElement(r_elem, f"{{{WML_NS}}}t")
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_elem.text = line if line else " "

    return comment


def _finalize_comments(output_path: str, comments_xml: bytes | None) -> None:
    """在 python-docx 保存后，将批注注入到 docx zip 包中。

    python-docx 不直接支持运行时添加 comments part，
    因此我们在保存后重新打开 zip，注入 comments.xml 并更新关系。
    """
    if not comments_xml:
        return

    import zipfile
    import shutil
    import os

    tmp_path = output_path + ".tmp"
    with zipfile.ZipFile(output_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == "word/document.xml":
                    # 在 document.xml 的关系中不需要额外操作，
                    # commentRangeStart/End 已经内联在 XML 中
                    pass

                elif item.filename == "[Content_Types].xml":
                    # 添加 comments 内容类型
                    ct_xml = etree.fromstring(data)
                    # 检查是否已存在
                    exists = ct_xml.find(
                        f'{{http://schemas.openxmlformats.org/package/2006/content-types}}Override[@PartName="/word/comments.xml"]'
                    )
                    if exists is None:
                        override = etree.SubElement(ct_xml, f"{{{CONTENT_TYPES_NS}}}Override")
                        override.set("PartName", "/word/comments.xml")
                        override.set(
                            "ContentType",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
                        )
                    data = etree.tostring(ct_xml, xml_declaration=True, encoding="UTF-8", standalone=True)

                elif item.filename == "word/_rels/document.xml.rels":
                    # 添加 comments 关系
                    rels_xml = etree.fromstring(data)
                    exists = rels_xml.find(
                        f'{{{REL_NS}}}Relationship[@Target="comments.xml"]'
                    )
                    if exists is None:
                        # 找到最大 rId
                        max_id = 0
                        for rel in rels_xml:
                            rid = rel.get("Id", "")
                            if rid.startswith("rId"):
                                try:
                                    max_id = max(max_id, int(rid[3:]))
                                except ValueError:
                                    pass
                        new_rid = f"rId{max_id + 1}"
                        rel_elem = etree.SubElement(rels_xml, f"{{{REL_NS}}}Relationship")
                        rel_elem.set("Id", new_rid)
                        rel_elem.set(
                            "Type",
                            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
                        )
                        rel_elem.set("Target", "comments.xml")
                    data = etree.tostring(rels_xml, xml_declaration=True, encoding="UTF-8", standalone=True)

                zout.writestr(item, data)

            # 添加 comments.xml
            zout.writestr("word/comments.xml", comments_xml)

    # 替换原文件
    shutil.move(tmp_path, output_path)


# ── 辅助函数 ──

def _extract_original_text(annotated_md: str, result: dict) -> str:
    """从 annotated_md 或评审结果中提取原文。"""
    # 方式1: 从 annotated_md 中解析段落
    if annotated_md:
        # 标准的 build_annotated_markdown 输出以 "# 审稿标注报告" 开头
        if "审稿标注报告" in annotated_md or "### 段落 " in annotated_md:
            paragraphs = []
            in_para = False
            for line in annotated_md.split("\n"):
                if line.startswith("### 段落 "):
                    in_para = True
                    continue
                if in_para:
                    # 遇到下一个标题 → 退出段落模式
                    if line.startswith("### ") or line.startswith("## ") or line == "---":
                        if line.startswith("### 段落 "):
                            in_para = True
                            continue
                        in_para = False
                        continue
                    if line.startswith("> "):  # 批注行
                        continue
                    if line.strip():
                        paragraphs.append(line.strip())
            if paragraphs:
                return "\n\n".join(paragraphs)

    # 方式2: 从 meta 中尝试获取原文（如果 API 未来支持）
    meta_text = result.get("meta", {}).get("original_text", "")
    if meta_text:
        return meta_text

    # 方式3: 如果 annotated_md 存在但不是标准格式，直接使用
    if annotated_md and not annotated_md.startswith("#"):
        return annotated_md

    return ""


def _get_emoji(sub: str) -> str:
    return {"FULL": "", "REVIEW": "⚡", "ESCALATE": "🔴"}.get(sub, "")


def _get_agent_color(agent: str) -> RGBColor:
    return {
        "A2": RGBColor(0x25, 0x63, 0xEB),
        "A3": RGBColor(0x16, 0xA3, 0x4A),
        "A4": RGBColor(0xEA, 0x58, 0x0C),
        "A5": RGBColor(0xDC, 0x26, 0x26),
    }.get(agent, RGBColor(0x64, 0x74, 0x8B))


def _add_legend_run(paragraph, text: str, color: RGBColor):
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = color


# ── 顶层导出接口（供 API/CLI 使用）──

def export_review_to_docx(result: dict, output_path: str | None = None) -> str:
    """导出评审结果为 .docx，自动处理 comments 注入。

    Args:
        result: 来自 /api/v1/review 的完整评审 JSON
        output_path: 完整输出路径。若为 None，自动生成到临时目录。

    Returns:
        .docx 文件路径
    """
    if output_path is None:
        output_dir = tempfile.mkdtemp()
        return export_annotated_docx(result, output_dir=output_dir)
    else:
        out_dir = str(Path(output_path).parent)
        tmp_path = export_annotated_docx(result, output_dir=out_dir)
        import shutil
        shutil.move(tmp_path, output_path)
        return output_path
