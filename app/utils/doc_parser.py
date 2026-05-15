from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class ParsedDocument:
    def __init__(self, text: str, word_count: int, metadata: dict | None = None):
        self.text = text
        self.word_count = word_count
        self.metadata = metadata or {}

    @property
    def paragraphs(self) -> list[str]:
        return [p.strip() for p in self.text.split("\n\n") if p.strip()]

    def get_text_chunk(self, max_words: int = 3000) -> str:
        """返回前 max_words 词的预览片段"""
        words = self.text.split()
        return " ".join(words[:max_words])


def parse_document(file_path: str | Path) -> ParsedDocument:
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        text = file_path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".md":
        text = file_path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".docx":
        text = _parse_docx(file_path)
    elif suffix == ".pdf":
        text = _parse_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Supported: .txt, .md, .docx, .pdf")

    word_count = len(text.split())
    return ParsedDocument(text=text, word_count=word_count, metadata={"filename": file_path.name, "format": suffix})


def _parse_docx(file_path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def _parse_pdf(file_path: Path) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")

    reader = PdfReader(str(file_path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)
